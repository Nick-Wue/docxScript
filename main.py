
import tkinter as tk
from tkinter import filedialog
from docx import Document
import docxedit

if __name__ == '__main__':

    root = tk.Tk()
    root.withdraw()

    path = filedialog.askopenfilename()
    doc = Document(path)
    # adds <lb/> to any linebreak
    docxedit.replace_string(doc, "\n", "<lb/>\n ")
    new_doc = Document()

    # iterate over all paragraphs and copy contents to new paragraphs
    for p_old in doc.paragraphs:
        para = new_doc.add_paragraph("<p>")
        '''new_doc.paragraphs[-1].paragraph_format.alignment = p_old.paragraph_format.alignment
        new_doc.paragraphs[-1].paragraph_format.first_line_indent = p_old.paragraph_format.first_line_indent
        new_doc.paragraphs[-1].paragraph_format.keep_together = p_old.paragraph_format.keep_together
        new_doc.paragraphs[-1].paragraph_format.keep_with_next = p_old.paragraph_format.keep_with_next
        new_doc.paragraphs[-1].paragraph_format.left_indent = p_old.paragraph_format.left_indent
        new_doc.paragraphs[-1].paragraph_format.line_spacing = p_old.paragraph_format.line_spacing
        new_doc.paragraphs[-1].paragraph_format.page_break_before = p_old.paragraph_format.page_break_before
        new_doc.paragraphs[-1].paragraph_format.right_indent = p_old.paragraph_format.right_indent
        new_doc.paragraphs[-1].paragraph_format.space_after = p_old.paragraph_format.space_after
        new_doc.paragraphs[-1].paragraph_format.space_before = p_old.paragraph_format.space_before
        new_doc.paragraphs[-1].paragraph_format.widow_control = p_old.paragraph_format.widow_control'''


        for r_old in p_old.runs:
            para.add_run(r_old.text)
            '''para.runs[-1].font.italic = r_old.font.italic
            para.runs[-1].font.all_caps = r_old.font.all_caps
            para.runs[-1].font.complex_script = r_old.font.complex_script
            para.runs[-1].font.double_strike = r_old.font.double_strike
            para.runs[-1].font.emboss = r_old.font.emboss
            para.runs[-1].font.hidden = r_old.font.hidden
            para.runs[-1].font.highlight_color = r_old.font.highlight_color
            para.runs[-1].font.imprint = r_old.font.imprint
            para.runs[-1].font.math = r_old.font.math
            para.runs[-1].font.name = r_old.font.name
            para.runs[-1].font.no_proof = r_old.font.no_proof
            para.runs[-1].font.outline = r_old.font.outline
            para.runs[-1].font.rtl = r_old.font.rtl
            para.runs[-1].font.shadow = r_old.font.shadow
            para.runs[-1].font.small_caps = r_old.font.small_caps
            para.runs[-1].font.snap_to_grid = r_old.font.snap_to_grid
            para.runs[-1].font.spec_vanish = r_old.font.spec_vanish
            para.runs[-1].font.strike = r_old.font.strike
            para.runs[-1].font.subscript = r_old.font.subscript
            para.runs[-1].font.superscript = r_old.font.superscript
            para.runs[-1].font.underline = r_old.font.underline
            para.runs[-1].font.web_hidden = r_old.font.web_hidden
            para.runs[-1].font.bold = r_old.font.bold
            para.runs[-1].font.size = r_old.font.size
            para.runs[-1].font.name = r_old.font.name
            para.runs[-1].font.cs_bold = r_old.font.cs_bold
            para.runs[-1].font.cs_italic = r_old.font.cs_italic'''



        para.add_run("</p>")

    new_doc.save("result.docx")

